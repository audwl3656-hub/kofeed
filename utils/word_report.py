"""결과보고서 Word(.docx) 생성 유틸리티."""
import io
import math
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.config import get_component_from_col, get_sample_from_col


# ── 헬퍼 ────────────────────────────────────────────────────
def _set_font(run, size=10, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def _para(doc, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def _heading(doc, text, level=1):
    sizes = {1: 15, 2: 13, 3: 11}
    _para(doc, text, size=sizes.get(level, 11), bold=True, space_before=8, space_after=6)


def _z_rgb(z):
    try:
        v = float(z)
        if math.isnan(v):
            return None
        if abs(v) <= 2:
            return (198, 239, 206)   # green
        elif abs(v) <= 3:
            return (255, 235, 156)   # yellow
        else:
            return (255, 199, 206)   # red
    except (TypeError, ValueError):
        return None


def _cell_bg(cell, rgb):
    """테이블 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_write(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(str(text) if text is not None else "")
    _set_font(run, size=size, bold=bold)
    if color_rgb:
        _cell_bg(cell, color_rgb)


def _add_table(doc, headers, rows, col_widths_cm=None, header_rgb=(68, 114, 196)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # 헤더
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        _cell_write(hdr_row.cells[i], h, bold=True, color_rgb=header_rgb)
        hdr_row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    # 데이터
    for row_data in rows:
        tr = table.add_row()
        for i, val in enumerate(row_data):
            cell_val = val if not isinstance(val, tuple) else val[0]
            z_fill   = val[1] if isinstance(val, tuple) else None
            _cell_write(tr.cells[i], cell_val,
                        align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                        color_rgb=z_fill)
    # 열 너비
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ── 결과보고서 Word ──────────────────────────────────────────
def generate_word_summary(
    df,
    z_all,
    z_method: dict,
    group_stats: dict,
    value_cols: list,
    inst_field: str,
    generated_at: str,
    samples: list,
    participant_map: dict = None,
    subtitle: str = "",
    period_배부: str = "",
    period_회신: str = "",
    period_보고서: str = "",
    sample_note: str = "",
    summary_text: str = "",
    sample_comp_text: dict = None,
    cfg=None,
) -> bytes:
    """결과보고서 Word 문서."""
    doc = Document()

    # 여백 설정 (2cm)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    _name_to_code = {v: k for k, v in (participant_map or {}).items()}
    raw_names  = df[inst_field].fillna("").astype(str).tolist()
    inst_codes = [_name_to_code.get(n, n) for n in raw_names]

    non_nir = [c for c in value_cols if group_stats.get(c)]
    seen_comps: list = []
    comp_to_cols: dict = {}
    for col in non_nir:
        comp = get_component_from_col(col, samples) or col
        if comp not in seen_comps:
            seen_comps.append(comp)
        comp_to_cols.setdefault(comp, []).append(col)

    # ── 표지 ──────────────────────────────────────────────
    _para(doc, "한국사료협회 비교분析 결과", size=22, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=60, space_after=8)
    if subtitle:
        _para(doc, subtitle, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _para(doc, generated_at[:10].replace("-", ".  ") + ".",
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "한국사료협회 사료기술연구소",
          size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # ── 1. 비교분析 개요 ──────────────────────────────────
    _heading(doc, "1. 비교분析 개요")

    _heading(doc, "가. 기간", level=2)
    _para(doc, f"1) 시료 배부 : {period_배부}", size=11)
    _para(doc, f"2) 분析 및 결과회신 : {period_회신}", size=11)
    _para(doc, f"3) 결과 통계처리 및 보고서 작성 : {period_보고서}", size=11, space_after=8)

    _heading(doc, "나. 시료 및 분析항목", level=2)

    # 사료별 성분 자동 계산
    _SECTION_GROUPS = {"아미노산", "NIR"}
    _comp_to_group: dict = {}
    if cfg is not None:
        from utils.config import get_component_groups as _get_cg, get_group_order as _get_go
        for _gn, _items in _get_cg(cfg).items():
            for _it in _items:
                _comp_to_group[_it["name"]] = _gn
        _group_order = _get_go(cfg)
        _enabled_sg = [g["name"] for g in _group_order if g["enabled"] and g["name"] in _SECTION_GROUPS]
    else:
        _enabled_sg = []

    sample_to_comps: dict = {}
    for comp in seen_comps:
        if _comp_to_group.get(comp) in _SECTION_GROUPS:
            continue
        for col in comp_to_cols.get(comp, []):
            s = get_sample_from_col(col, samples)
            if s and s in samples:
                sample_to_comps.setdefault(s, [])
                if comp not in sample_to_comps[s]:
                    sample_to_comps[s].append(comp)

    ov_rows = []
    for i, s in enumerate(samples):
        comps_s = list(sample_to_comps.get(s, []))
        for sg in _enabled_sg:
            comps_s.append(sg)
        if not comps_s:
            continue
        comp_text = (sample_comp_text or {}).get(s, "").strip() or ", ".join(comps_s)
        ov_rows.append([f"{s}(샘플{i+1})", comp_text])

    _add_table(doc, ["시료", "분析항목"], ov_rows, col_widths_cm=[4, 14])
    if sample_note:
        _para(doc, f"* {sample_note}", size=9, color=(128, 128, 128), space_after=8)

    _heading(doc, "다. 통계처리방법", level=2)
    _para(doc, "1) 평균(x-bar), 표준편차(σ), 변이계수(CV) 등을 산출함.", size=11)
    _para(doc, "2) 시험소 간 비교숙련도 시험용 Robust Z-score, Outlier : KS Q ISO 13528에 따라 산출함.", size=11, space_after=8)

    doc.add_page_break()

    # ── 2. 통계 요약 ──────────────────────────────────────
    _heading(doc, "2. 통계 요약")
    stat_rows = []
    for comp in seen_comps:
        for col in comp_to_cols.get(comp, []):
            st = group_stats.get(col, {})
            if not st:
                continue
            s_label = get_sample_from_col(col, samples) or "-"
            vals = pd.to_numeric(df[col], errors="coerce").dropna()
            cv = st.get("cv", float("nan"))
            stat_rows.append([
                comp, s_label,
                str(int(st.get("n", len(vals)))),
                f"{st.get('mean', 0):.4f}",
                f"{st.get('median', 0):.4f}",
                f"{st.get('std', 0):.4f}",
                f"{cv:.2f}" if not math.isnan(cv) else "N/A",
                f"{vals.min():.4f}" if not vals.empty else "-",
                f"{vals.max():.4f}" if not vals.empty else "-",
            ])
    _add_table(doc, ["성분", "사료", "n", "평균", "중앙값", "표준편차", "CV(%)", "최솟값", "최댓값"],
               stat_rows, col_widths_cm=[3, 3, 1.2, 2.2, 2.2, 2.2, 1.8, 2.2, 2.2])

    doc.add_page_break()

    # ── 3. Z-score 결과 ───────────────────────────────────
    _heading(doc, "3. Z-score 결과")
    _heading(doc, "가. 전체 Z-score", level=2)

    for comp in seen_comps:
        _para(doc, f"▶ {comp}", size=10, bold=True, space_before=6, space_after=3)
        cols_c = comp_to_cols.get(comp, [])
        z_hdrs = ["기관코드"] + [f"{get_sample_from_col(c, samples) or c}_값" for c in cols_c] + \
                               [f"{get_sample_from_col(c, samples) or c}_Z" for c in cols_c]
        z_rows = []
        for ri, (idx, row) in enumerate(df.iterrows()):
            r_data = [inst_codes[ri]]
            for col in cols_c:
                v = row.get(col, "")
                try:
                    r_data.append(f"{float(v):.4f}" if v != "" and not pd.isna(v) else "-")
                except (TypeError, ValueError):
                    r_data.append("-")
            for col in cols_c:
                z = z_all.loc[idx, col]
                try:
                    z_v = round(float(z), 3) if not pd.isna(z) else None
                except (TypeError, ValueError):
                    z_v = None
                rgb = _z_rgb(z_v)
                r_data.append((f"{z_v:.3f}" if z_v is not None else "N/A", rgb))
            z_rows.append(r_data)
        _add_table(doc, z_hdrs, z_rows)

    doc.add_paragraph()
    _heading(doc, "나. 방법별 Z-score", level=2)

    for comp in seen_comps:
        _para(doc, f"▶ {comp}", size=10, bold=True, space_before=6, space_after=3)
        cols_c = comp_to_cols.get(comp, [])
        z_hdrs = ["기관코드"] + [f"{get_sample_from_col(c, samples) or c}_값" for c in cols_c] + \
                               [f"{get_sample_from_col(c, samples) or c}_Z방법별" for c in cols_c]
        z_rows = []
        for ri, (idx, row) in enumerate(df.iterrows()):
            r_data = [inst_codes[ri]]
            for col in cols_c:
                v = row.get(col, "")
                try:
                    r_data.append(f"{float(v):.4f}" if v != "" and not pd.isna(v) else "-")
                except (TypeError, ValueError):
                    r_data.append("-")
            for col in cols_c:
                zm = z_method.get(col, pd.Series())
                try:
                    z_v = round(float(zm.loc[idx]), 3) if not pd.isna(zm.loc[idx]) else None
                except (KeyError, TypeError, ValueError):
                    z_v = None
                rgb = _z_rgb(z_v)
                r_data.append((f"{z_v:.3f}" if z_v is not None else "N/A", rgb))
            z_rows.append(r_data)
        _add_table(doc, z_hdrs, z_rows)

    # ── 4. 분析결과 요약 ──────────────────────────────────
    if summary_text:
        doc.add_page_break()
        _heading(doc, "4. 분析결과 요약")
        for line in summary_text.splitlines():
            _para(doc, line, size=11, space_after=3)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
