"""
전체요약 보고서 DOCX 생성.
차트는 matplotlib으로 PNG 렌더링 후 삽입.
표, 제목, 텍스트는 python-docx로 작성.
"""
import io
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.config import (
    get_component_from_col, get_sample_from_col, get_col_suffix, get_base_col,
    get_method_options,
)


# ── 한글 폰트 설정 ──────────────────────────────────────────────
def _get_ko_font():
    """시스템에서 사용 가능한 한글 폰트 경로 반환."""
    candidates = [
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
        "C:/Windows/Fonts/malgun.ttf",
        "C:/Windows/Fonts/gulim.ttc",
    ]
    for p in candidates:
        try:
            fm.FontProperties(fname=p)
            return p
        except Exception:
            continue
    return None


_KO_FONT_PATH = _get_ko_font()
_KO_FONT_NAME = "NanumGothic"

def _setup_mpl_font():
    if _KO_FONT_PATH:
        fe = fm.FontEntry(fname=_KO_FONT_PATH, name=_KO_FONT_NAME)
        fm.fontManager.ttflist.insert(0, fe)
        plt.rcParams["font.family"] = _KO_FONT_NAME
    plt.rcParams["axes.unicode_minus"] = False


# ── 스타일 헬퍼 ─────────────────────────────────────────────────
_HDR_COLOR = RGBColor(0x44, 0x72, 0xC4)   # #4472C4
_ALT_COLOR  = RGBColor(0xF0, 0xF4, 0xFA)  # #f0f4fa
_GRID_COLOR = RGBColor(0x00, 0x00, 0x00)

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}")
    tcPr.append(shd)


def _set_cell_border(cell, color="000000", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=10, bold=False,
              color: RGBColor = None, font_name="맑은 고딕"):
    para.alignment = align
    for run in para.runs:
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.name = font_name
        if color:
            run.font.color.rgb = color
    return para


def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_paragraph(doc: Document, text: str, size_pt=10, bold=False,
                   indent_cm=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.name  = "맑은 고딕"
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    return p


# ── 차트 → PNG bytes ────────────────────────────────────────────
def _cv_chart_png(cv_data: dict, seen_comps: list, valid_stat_samples: list,
                  dpi=150) -> bytes:
    """각 성분별 변이계수(CV%) 가로 막대 차트 → PNG bytes."""
    _setup_mpl_font()
    comps = [c for c in seen_comps if c in cv_data]
    feeds = [s for s in valid_stat_samples
             if any(s in cv_data.get(c, {}) for c in comps)]
    if not comps or not feeds:
        return b""

    palette = ["#2563eb", "#d97706", "#16a34a", "#dc2626",
               "#7c3aed", "#0891b2", "#be185d"]
    n_comps = len(comps)
    n_feeds = len(feeds)
    x = np.arange(n_comps)
    bar_w = 0.6 / n_feeds

    fig, ax = plt.subplots(figsize=(max(8, n_comps * 0.6), 4))
    for fi, feed in enumerate(feeds):
        vals = [cv_data.get(c, {}).get(feed, np.nan) for c in comps]
        offset = (fi - (n_feeds - 1) / 2) * bar_w
        ax.bar(x + offset, vals, width=bar_w * 0.85,
               color=palette[fi % len(palette)], label=feed)

    ax.set_xticks(x)
    ax.set_xticklabels([c[:8] for c in comps], fontsize=8, rotation=30, ha="right")
    ax.set_ylabel("CV (%)", fontsize=9)
    ax.set_title("각 성분별 변이계수(CV%)", fontsize=11)
    ax.legend(fontsize=8, loc="upper right")
    ax.yaxis.grid(True, linestyle="--", alpha=0.5)
    ax.set_axisbelow(True)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi)
    plt.close(fig)
    return buf.getvalue()


def _zscore_chart_png(comp: str, z_src, df, inst_names: list, idx_list: list,
                      samples: list, non_nir_set: set,
                      group_by_method=False, min_n=1, cfg=None, dpi=120) -> list:
    """단일 성분 Z-score 막대그래프 → [(title, PNG bytes), ...]"""
    _setup_mpl_font()
    import re as _re

    sfx_list = [""]
    for n in range(2, 10):
        if any(f"{comp}_{s}_{n}" in non_nir_set for s in samples):
            sfx_list.append(f"_{n}")
        else:
            break

    valid_samps = [s for s in samples
                   if any(f"{comp}_{s}{sfx}" in non_nir_set for sfx in sfx_list)]
    if not valid_samps:
        return []

    is_fat = "조지방" in str(comp)

    def _sol_abbr(val):
        m = str(val).upper()
        if "PETROLEUM" in m or "석유에테르" in val or "석유" in val: return "P"
        if "DIETHYL" in m: return "DE"
        if "헥산" in val or "HEXAN" in m: return "H"
        if "에탄올" in val or "ETHANOL" in m: return "EtOH"
        if "아세톤" in val or "ACETON" in m: return "Ac"
        if "ETHER" in m or "에테르" in val: return "E"
        return ""

    def _get_zv(row_idx, col):
        try:
            import pandas as _pd
            if isinstance(z_src, _pd.DataFrame):
                return z_src.at[row_idx, col] if col in z_src.columns else np.nan
            ser = z_src.get(col)
            return ser.at[row_idx] if ser is not None else np.nan
        except Exception:
            return np.nan

    results = []
    for s in valid_samps:
        grouped: dict = {}
        for sfx in sfx_list:
            col = f"{comp}_{s}{sfx}"
            if col not in non_nir_set:
                continue
            mc = f"{comp}_방법{sfx}"
            for inst, row_idx in zip(inst_names, idx_list):
                zv = _get_zv(row_idx, col)
                try:
                    zv = float(zv)
                except Exception:
                    continue
                if np.isnan(zv):
                    continue
                meth_str = ""
                if mc in df.columns:
                    try:
                        meth_str = str(df.at[row_idx, mc]).strip()
                    except Exception:
                        pass
                if not meth_str or meth_str == "nan":
                    meth_str = "(방법 미기재)"
                grp_key = meth_str if group_by_method else "__ALL__"
                if is_fat:
                    sc = f"{comp}_용매{sfx}"
                    sol_val = ""
                    if sc in df.columns:
                        try:
                            v = str(df.at[row_idx, sc]).strip()
                            sol_val = "" if v in ("nan", "해당없음", "-", "") else v
                        except Exception:
                            pass
                    if not sol_val and sfx:
                        sc0 = f"{comp}_용매"
                        if sc0 in df.columns:
                            try:
                                v = str(df.at[row_idx, sc0]).strip()
                                sol_val = "" if v in ("nan", "해당없음", "-", "") else v
                            except Exception:
                                pass
                    abbr = _sol_abbr(sol_val) if sol_val else _sol_abbr(meth_str)
                    label = f"{inst}-{abbr}" if abbr else str(inst)
                else:
                    label = str(inst)
                grouped.setdefault(grp_key, []).append((label, zv))

        grouped = {k: v for k, v in grouped.items() if len(v) >= min_n}
        if not grouped:
            continue

        ordered_methods = get_method_options(cfg, comp=comp) if cfg is not None else []
        def _gk2(k, _om=ordered_methods):
            try: return _om.index(k)
            except: return len(_om)

        for grp_key, bars in sorted(grouped.items(), key=lambda x: _gk2(x[0])):
            bars.sort(key=lambda item: item[1])
            labels = [b[0] for b in bars]
            z_vals = [b[1] for b in bars]
            bar_colors = ["#dc2626" if abs(z) >= 3 else "#4472C4" for z in z_vals]

            import math as _m
            abs_max = max(abs(min(z_vals)), abs(max(z_vals)), 2.5)
            y_hi = float(_m.ceil(abs_max * 1.15))

            if y_hi <= 5: tick_step = 1
            elif y_hi <= 10: tick_step = 2
            else: tick_step = 5

            fig_w = max(6, len(bars) * 0.35)
            fig, ax = plt.subplots(figsize=(fig_w, 3.5))
            x = np.arange(len(bars))
            ax.bar(x, z_vals, color=bar_colors, width=0.5)
            ax.axhline(0, color="#888888", linewidth=0.8)
            ax.set_xticks(x)
            ax.set_xticklabels(labels, fontsize=6, rotation=45, ha="right")
            ax.set_ylim(-y_hi, y_hi)
            ax.yaxis.set_ticks(np.arange(
                int(-y_hi // tick_step) * tick_step,
                int(y_hi // tick_step) * tick_step + tick_step,
                tick_step,
            ))
            ax.yaxis.grid(True, linestyle="--", alpha=0.4)
            ax.set_axisbelow(True)
            title = (f"{s}({comp}-{grp_key})"
                     if group_by_method and grp_key not in ("__ALL__", "(방법 미기재)")
                     else f"{s}({comp})")
            ax.set_title(title, fontsize=9)
            for xi, zv in enumerate(z_vals):
                ax.text(xi, zv + (0.05 if zv >= 0 else -0.15), f"{zv:.2f}",
                        ha="center", va="bottom" if zv >= 0 else "top",
                        fontsize=5.5, color="#dc2626" if abs(zv) >= 3 else "#444444")
            fig.tight_layout()
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=dpi)
            plt.close(fig)
            results.append((title, buf.getvalue()))

    return results


# ── 표 스타일 적용 ──────────────────────────────────────────────
def _style_header_row(table, row_idx=0):
    """헤더 행 배경 파랑, 텍스트 굵게."""
    row = table.rows[row_idx]
    for cell in row.cells:
        _set_cell_bg(cell, _HDR_COLOR)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.name = "맑은 고딕"
                run.font.color.rgb = RGBColor(0, 0, 0)


def _style_data_rows(table, start=1):
    for ri, row in enumerate(table.rows[start:], start=start):
        bg = _ALT_COLOR if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
        for cell in row.cells:
            _set_cell_bg(cell, bg)
            _set_cell_border(cell)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.name = "맑은 고딕"


def _add_image_to_doc(doc: Document, img_bytes: bytes, width_cm=15):
    if not img_bytes:
        return
    buf = io.BytesIO(img_bytes)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))


# ── 메인 생성 함수 ──────────────────────────────────────────────
def generate_docx_summary(
    df,
    z_all,
    z_method: dict,
    group_stats: dict,
    value_cols: list,
    inst_field: str,
    generated_at: str = None,
    samples: list = None,
    participant_map: dict = None,
    subtitle: str = "",
    period_배부: str = "",
    period_회신: str = "",
    period_보고서: str = "",
    sample_note: str = "",
    summary_text: str = "",
    cfg=None,
) -> bytes:
    import pandas as pd

    if samples is None:
        from utils.config import get_samples
        samples = get_samples()
    generated_at = generated_at or datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── 데이터 준비 ──
    non_nir = [c for c in value_cols if group_stats.get(c)]
    non_nir_set = set(non_nir)
    seen_comps: list = []
    comp_to_cols: dict = {}
    for col in non_nir:
        comp = get_component_from_col(col, samples) or col
        if comp not in seen_comps:
            seen_comps.append(comp)
        comp_to_cols.setdefault(comp, []).append(col)

    valid_stat_samples = [s for s in samples
                          if any(f"{c}_{s}" in non_nir_set for c in seen_comps)]
    _name_to_code = {v: k for k, v in (participant_map or {}).items()}
    raw_inst_names = df[inst_field].fillna("").astype(str).tolist()
    inst_names = [_name_to_code.get(n, n) for n in raw_inst_names]
    idx_list = df.index.tolist()

    def fmt(v, dec=2):
        try:
            f = float(v)
            return "-" if np.isnan(f) else f"{f:.{dec}f}"
        except Exception:
            return "-" if str(v).strip() == "" else str(v)

    # ── 문서 생성 ──
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    section.left_margin   = Mm(20)
    section.right_margin  = Mm(20)
    section.top_margin    = Mm(25)
    section.bottom_margin = Mm(20)

    # ── 표지 ──
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("한국사료협회 비교분석 결과")
    run.font.name = "맑은 고딕"
    run.font.size = Pt(24)
    run.font.bold = True

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub_p.add_run(subtitle)
        r.font.name = "맑은 고딕"
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(period_보고서 or generated_at)
    r.font.name = "맑은 고딕"
    r.font.size = Pt(12)

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org_p.add_run("한국사료협회 사료기술연구소")
    r.font.name = "맑은 고딕"
    r.font.size = Pt(12)

    doc.add_page_break()

    # ── 1. 비교분석 개요 ──
    _add_heading(doc, "1. 비교분석 개요", 1)
    _add_heading(doc, "가. 기간", 2)
    _add_paragraph(doc, f"1) 시료 배부 : {period_배부}")
    _add_paragraph(doc, f"2) 분석 및 결과회신 : {period_회신}")
    _add_paragraph(doc, f"3) 보고서 발송 : {period_보고서}")

    _add_heading(doc, "나. 시료 및 분석항목", 2)
    sample_to_comps: dict = {}
    for comp in seen_comps:
        for col in comp_to_cols.get(comp, []):
            s = get_sample_from_col(col, samples)
            if s:
                sample_to_comps.setdefault(s, set()).add(comp)
    for s, comps_set in sample_to_comps.items():
        ordered = [c for c in seen_comps if c in comps_set]
        _add_paragraph(doc, f"- {s}: {', '.join(ordered)}")
    if sample_note:
        _add_paragraph(doc, f"※ {sample_note}", size_pt=9)

    _add_heading(doc, "다. 통계처리방법", 2)
    _add_paragraph(doc, "1) 평균(x̄), 표준편차(σ), 변이계수(CV) 등을 산출함.")
    _add_paragraph(doc, "2) Robust Z-score: Z = (x - Median) / ((Q3-Q1) × 0.7413)")
    _add_paragraph(doc, "   |Z| ≤ 2: 적합   2 < |Z| ≤ 3: 경고   |Z| > 3: 부적합")

    _add_heading(doc, "라. 참가회원사", 2)
    _inst_real = sorted(set(raw_inst_names))
    _add_paragraph(doc, ", ".join(_inst_real))

    doc.add_page_break()

    # ── 2. 비교분석 결과 ──
    _add_heading(doc, "2. 비교분석 결과", 1)

    # ── 가. 분석결과와 통계 ──
    _add_heading(doc, "가. 분석결과와 통계", 2)

    def _is_dead(comp, s):
        for n in range(1, 10):
            sfx = "" if n == 1 else f"_{n}"
            if f"{comp}_{s}{sfx}" in non_nir_set:
                return False
        return True

    for comp in seen_comps:
        _add_paragraph(doc, comp, bold=True, size_pt=10, space_before=6)

        method_entries = []
        for n in range(1, 10):
            sfx = "" if n == 1 else f"_{n}"
            mc = f"{comp}_방법{sfx}"
            if mc not in df.columns:
                if n > 1: break
                continue
            for m in df[mc].dropna().astype(str).unique():
                if m.strip():
                    method_entries.append((mc, sfx, m.strip()))
        if not method_entries:
            continue

        vs = [s for s in valid_stat_samples if not _is_dead(comp, s)]
        if not vs:
            continue

        n_col = 4  # N, 평균, 표준편차, 변이계수
        headers = ["분석법"] + [f"{s}\n(N/평균/SD/CV%)" for s in vs]
        tbl = doc.add_table(rows=1, cols=1 + len(vs) * n_col)
        tbl.style = "Table Grid"

        # 헤더: 분석법 | (사료 x 4cols)
        hdr_row = tbl.rows[0]
        hdr_row.cells[0].text = "분석법"
        for si, s in enumerate(vs):
            c0 = 1 + si * n_col
            hdr_row.cells[c0].text = s
            for offset in range(1, n_col):
                hdr_row.cells[c0 + offset].text = ""
        _style_header_row(tbl, 0)

        # 소헤더: N | 평균 | SD | CV%
        sub_row = tbl.add_row()
        sub_row.cells[0].text = ""
        for si in range(len(vs)):
            c0 = 1 + si * n_col
            for j, h in enumerate(["N", "평균", "SD", "CV%"]):
                sub_row.cells[c0 + j].text = h
        _style_header_row(tbl, 1)

        # 데이터 행
        row_idx_tbl = 2
        for mc, sfx, meth in method_entries:
            dr = tbl.add_row()
            dr.cells[0].text = meth
            for si, s in enumerate(vs):
                col = f"{comp}_{s}{sfx}"
                c0 = 1 + si * n_col
                if col not in df.columns:
                    for j in range(n_col): dr.cells[c0+j].text = "-"
                    continue
                mask = (df[mc].fillna("").astype(str).str.strip() == meth)
                vals = pd.to_numeric(df.loc[mask, col], errors="coerce").dropna()
                if len(vals) == 0:
                    for j in range(n_col): dr.cells[c0+j].text = "-"
                    continue
                mean_ = vals.mean()
                std_  = vals.std(ddof=1) if len(vals) > 1 else float("nan")
                cv_   = (std_/mean_*100) if mean_ != 0 and not np.isnan(std_) else float("nan")
                dr.cells[c0+0].text = str(len(vals))
                dr.cells[c0+1].text = fmt(mean_)
                dr.cells[c0+2].text = fmt(std_)
                dr.cells[c0+3].text = fmt(cv_, 1)
            row_idx_tbl += 1

        # 전체 행
        wr = tbl.add_row()
        wr.cells[0].text = f"{comp} 전체"
        for si, s in enumerate(vs):
            c0 = 1 + si * n_col
            all_vals = []
            for n2 in range(1, 10):
                sx2 = "" if n2 == 1 else f"_{n2}"
                c2 = f"{comp}_{s}{sx2}"
                if c2 not in df.columns:
                    if n2 > 1: break
                    continue
                all_vals.extend(pd.to_numeric(df[c2], errors="coerce").dropna().tolist())
            if not all_vals:
                for j in range(n_col): wr.cells[c0+j].text = "-"
                continue
            vals = pd.Series(all_vals)
            mean_ = vals.mean()
            std_  = vals.std(ddof=1) if len(vals) > 1 else float("nan")
            cv_   = (std_/mean_*100) if mean_ != 0 and not np.isnan(std_) else float("nan")
            wr.cells[c0+0].text = str(len(vals))
            wr.cells[c0+1].text = fmt(mean_)
            wr.cells[c0+2].text = fmt(std_)
            wr.cells[c0+3].text = fmt(cv_, 1)

        _style_data_rows(tbl, start=2)

    doc.add_page_break()

    # ── 나. 분석결과 요약 (CV 차트) ──
    _add_heading(doc, "나. 분석결과 요약", 2)
    cv_data: dict = {}
    for comp in seen_comps:
        for s in valid_stat_samples:
            col = f"{comp}_{s}"
            try:
                cv_f = float(group_stats.get(col, {}).get("cv", float("nan")))
            except Exception:
                cv_f = float("nan")
            cv_data.setdefault(comp, {})[s] = cv_f

    cv_png = _cv_chart_png(cv_data, seen_comps, valid_stat_samples)
    if cv_png:
        _add_image_to_doc(doc, cv_png, width_cm=15)

    # CV 표
    if cv_data:
        comps_with_data = [c for c in seen_comps if c in cv_data]
        feeds_with_data = [s for s in valid_stat_samples
                           if any(s in cv_data.get(c, {}) for c in comps_with_data)]
        if comps_with_data and feeds_with_data:
            cv_tbl = doc.add_table(rows=1 + len(feeds_with_data),
                                   cols=1 + len(comps_with_data))
            cv_tbl.style = "Table Grid"
            cv_tbl.rows[0].cells[0].text = "사료종류"
            for ci, comp in enumerate(comps_with_data):
                cv_tbl.rows[0].cells[ci+1].text = comp
            _style_header_row(cv_tbl, 0)
            for fi, feed in enumerate(feeds_with_data):
                row = cv_tbl.rows[fi+1]
                row.cells[0].text = feed
                for ci, comp in enumerate(comps_with_data):
                    v = cv_data.get(comp, {}).get(feed)
                    row.cells[ci+1].text = (
                        f"{v:.2f}" if v is not None and not np.isnan(v) else "-"
                    )
            _style_data_rows(cv_tbl, start=1)

    if summary_text:
        doc.add_paragraph()
        for line in summary_text.splitlines():
            if line.strip():
                _add_paragraph(doc, line.strip())

    doc.add_page_break()

    # ── 다. 시료, 성분별 Z-score ──
    _add_heading(doc, "다. 시료, 성분별 Robust Z-score", 2)
    _add_zscore_section(doc, seen_comps, non_nir_set, samples, valid_stat_samples,
                        z_all, df, inst_names, idx_list, group_stats, fmt,
                        group_by_method=False, min_n=1, cfg=cfg)

    doc.add_page_break()

    # ── 라. 방법별 Z-score ──
    _add_heading(doc, "라. 방법별 Robust Z-score", 2)
    _add_zscore_section(doc, seen_comps, non_nir_set, samples, valid_stat_samples,
                        z_method, df, inst_names, idx_list, group_stats, fmt,
                        group_by_method=True, min_n=5, split_at=5, cfg=cfg)

    doc.add_page_break()

    # ── 판정 기준 ──
    _add_heading(doc, "판정 기준 (Robust Z-score)", 2)
    _add_paragraph(doc, "적합: |Z| ≤ 2.0   경고: 2.0 < |Z| ≤ 3.0   부적합: |Z| > 3.0")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_zscore_section(doc, seen_comps, non_nir_set, samples, valid_stat_samples,
                        z_src, df, inst_names, idx_list, group_stats, fmt,
                        group_by_method=False, min_n=1, split_at=0, cfg=None):
    import pandas as pd

    def _get_zv(row_idx, col):
        try:
            if isinstance(z_src, pd.DataFrame):
                return z_src.at[row_idx, col] if col in z_src.columns else np.nan
            ser = z_src.get(col)
            return ser.at[row_idx] if ser is not None else np.nan
        except Exception:
            return np.nan

    def _lab_sort_key(item):
        import re as _re
        txt = str(item[0])
        m = _re.match(r'[A-Za-z]*(\d+)', txt)
        return (0, int(m.group(1))) if m else (1, txt)

    for num, comp in enumerate(seen_comps, 1):
        sfx_list = [""]
        for n in range(2, 10):
            if any(f"{comp}_{s}_{n}" in non_nir_set for s in samples):
                sfx_list.append(f"_{n}")
            else:
                break

        valid_samples = [s for s in samples
                         if any(f"{comp}_{s}{sfx}" in non_nir_set for sfx in sfx_list)]
        if not valid_samples:
            continue

        # 방법별 그룹핑
        method_to_rows: dict = {}
        method_raw_vals: dict = {}
        for sfx in sfx_list:
            mc = f"{comp}_방법{sfx}"
            for inst, row_idx in zip(inst_names, idx_list):
                meth = str(df.at[row_idx, mc] if mc in df.columns else "").strip()
                if not meth:
                    meth = "(방법 미기재)"
                has_any = False
                row_data_vals = []
                for s in valid_samples:
                    col = f"{comp}_{s}{sfx}"
                    if col not in non_nir_set:
                        row_data_vals.append(("-", "-")); continue
                    raw = df.at[row_idx, col] if col in df.columns else ""
                    try:
                        rv = float(raw)
                        res_str = fmt(rv) if not np.isnan(rv) else "-"
                        if not np.isnan(rv):
                            has_any = True
                            method_raw_vals.setdefault(meth, {}).setdefault(s, []).append(rv)
                    except Exception:
                        res_str = "-"
                    zv = _get_zv(row_idx, col)
                    try:
                        zv_f = float(zv)
                        z_str = f"{zv_f:.2f}" if not np.isnan(zv_f) else "-"
                    except Exception:
                        z_str = "-"
                    row_data_vals.append((res_str, z_str))
                if not has_any:
                    continue
                lab_label = inst if sfx == "" else f"{inst}{sfx}"
                method_to_rows.setdefault(meth, []).append((lab_label, row_data_vals))

        if min_n > 1:
            method_to_rows = {m: r for m, r in method_to_rows.items() if len(r) >= min_n}
        if not method_to_rows:
            continue

        ordered_methods = get_method_options(cfg, comp=comp) if cfg is not None else []
        def _meth_order(item):
            try: return ordered_methods.index(item[0])
            except ValueError: return len(ordered_methods)

        sorted_methods = sorted(method_to_rows.items(), key=_meth_order)
        total_data_rows = sum(len(r) for _, r in sorted_methods)
        do_split = split_at > 0 and total_data_rows > split_at

        _add_paragraph(doc, f"{num}) {comp}", bold=True, size_pt=10, space_before=6)

        # 표 생성
        n_samp = len(valid_samples)
        # 열: 분석방법 | Lab | (결과, Z-score) x n_samp
        n_cols = 2 + n_samp * 2

        def _make_zscore_table(meth_rows_list, meth_name=None, median_override=None):
            tbl = doc.add_table(rows=2, cols=n_cols)
            tbl.style = "Table Grid"

            # 헤더 1행
            tbl.rows[0].cells[0].text = "분석방법"
            tbl.rows[0].cells[1].text = "Lab"
            for si, s in enumerate(valid_samples):
                c0 = 2 + si * 2
                if median_override and s in median_override:
                    med_str = f"{median_override[s]:.2f}"
                else:
                    col0 = f"{comp}_{s}"
                    gs = group_stats.get(col0, {})
                    try:
                        med = float(gs.get("median", float("nan")))
                        med_str = f"{med:.2f}" if not np.isnan(med) else "-"
                    except Exception:
                        med_str = "-"
                tbl.rows[0].cells[c0].text = f"{s} (중간값: {med_str})"
                tbl.rows[0].cells[c0+1].text = ""
            _style_header_row(tbl, 0)

            # 헤더 2행
            tbl.rows[1].cells[0].text = ""
            tbl.rows[1].cells[1].text = ""
            for si in range(n_samp):
                c0 = 2 + si * 2
                tbl.rows[1].cells[c0].text = "결과"
                tbl.rows[1].cells[c0+1].text = "Z-score"
            _style_header_row(tbl, 1)

            # 데이터 행
            sorted_rows = sorted(meth_rows_list, key=_lab_sort_key)
            first = True
            for lab, row_vals in sorted_rows:
                dr = tbl.add_row()
                dr.cells[0].text = meth_name if (first and meth_name) else ""
                dr.cells[1].text = lab
                for si, (res_str, z_str) in enumerate(row_vals):
                    c0 = 2 + si * 2
                    dr.cells[c0].text = res_str
                    # z-score 색상 (빨강 if |z|>3)
                    try:
                        zf = float(z_str)
                        dr.cells[c0+1].text = z_str
                        if abs(zf) > 3:
                            for para in dr.cells[c0+1].paragraphs:
                                for run in para.runs:
                                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                                    run.font.bold = True
                    except Exception:
                        dr.cells[c0+1].text = z_str
                first = False
            _style_data_rows(tbl, start=2)
            return tbl

        if do_split:
            for meth, mrows in sorted_methods:
                med_ov = {}
                for s in valid_samples:
                    m_vals = method_raw_vals.get(meth, {}).get(s, [])
                    if m_vals:
                        med_ov[s] = float(np.median(m_vals))
                _add_paragraph(doc, f"▶ {meth}", size_pt=9, space_before=3)
                _make_zscore_table(mrows, meth_name=meth, median_override=med_ov)
                doc.add_paragraph()
        else:
            all_rows = []
            for meth, mrows in sorted_methods:
                for lab, row_vals in mrows:
                    all_rows.append((lab, row_vals, meth))
            # 단일 표로
            tbl = doc.add_table(rows=2, cols=n_cols)
            tbl.style = "Table Grid"
            tbl.rows[0].cells[0].text = "분석방법"
            tbl.rows[0].cells[1].text = "Lab"
            for si, s in enumerate(valid_samples):
                c0 = 2 + si * 2
                col0 = f"{comp}_{s}"
                gs = group_stats.get(col0, {})
                try:
                    med = float(gs.get("median", float("nan")))
                    med_str = f"{med:.2f}" if not np.isnan(med) else "-"
                except Exception:
                    med_str = "-"
                tbl.rows[0].cells[c0].text = f"{s} (중간값: {med_str})"
                tbl.rows[0].cells[c0+1].text = ""
            _style_header_row(tbl, 0)
            tbl.rows[1].cells[0].text = ""
            tbl.rows[1].cells[1].text = ""
            for si in range(n_samp):
                c0 = 2 + si * 2
                tbl.rows[1].cells[c0].text = "결과"
                tbl.rows[1].cells[c0+1].text = "Z-score"
            _style_header_row(tbl, 1)

            prev_meth = None
            sorted_flat = []
            for meth, mrows in sorted_methods:
                for lab, row_vals in sorted(mrows, key=_lab_sort_key):
                    sorted_flat.append((lab, row_vals, meth))

            for lab, row_vals, meth in sorted_flat:
                dr = tbl.add_row()
                dr.cells[0].text = meth if meth != prev_meth else ""
                prev_meth = meth
                dr.cells[1].text = lab
                for si, (res_str, z_str) in enumerate(row_vals):
                    c0 = 2 + si * 2
                    dr.cells[c0].text = res_str
                    try:
                        zf = float(z_str)
                        dr.cells[c0+1].text = z_str
                        if abs(zf) > 3:
                            for para in dr.cells[c0+1].paragraphs:
                                for run in para.runs:
                                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                                    run.font.bold = True
                    except Exception:
                        dr.cells[c0+1].text = z_str
            _style_data_rows(tbl, start=2)
            doc.add_paragraph()

        # Z-score 그래프 삽입
        chart_list = _zscore_chart_png(
            comp, z_src, df, inst_names, idx_list,
            samples, non_nir_set,
            group_by_method=group_by_method, min_n=min_n, cfg=cfg,
        )
        for title, png_bytes in chart_list:
            _add_image_to_doc(doc, png_bytes, width_cm=14)
